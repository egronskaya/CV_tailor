import os
from docx import Document
from docx.shared import Pt, Inches
import yaml
from fpdf import FPDF
from typing import Dict, List, Tuple
import subprocess

class DocumentMaker:
    def __init__(self):
        # Load style guide
        style_guide_path = os.getenv('USER_STYLE_PATH', 'user_data/style/user_style.yaml')
        with open(style_guide_path, 'r') as f:
            self.style_guide = yaml.safe_load(f)
            
        # Document settings
        self.cv_format = os.getenv('CV_FORMAT', 'pdf')
        
        # Output settings
        self.output_dir = os.getenv('OUTPUT_DIR', 'output')
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            
        self.max_file_size = int(os.getenv('MAX_FILE_SIZE', '5')) * 1024 * 1024  # Convert MB to bytes
    
    def create_cv_documents(self, cv_data: Dict) -> Tuple[bytes, bytes]:
        """Create both DOCX and PDF versions of the CV."""
        # Create DOCX version
        docx_data = self._create_cv_docx(cv_data)
        
        # Create PDF version using LaTeX
        pdf_data = self._create_cv_pdf(cv_data)
        
        return docx_data, pdf_data
    
    def create_letter_documents(self, letters: List[Dict]) -> List[Tuple[bytes, bytes]]:
        """Create DOCX and PDF versions of each cover letter."""
        documents = []
        
        for letter in letters:
            docx_data = self._create_letter_docx(letter)
            pdf_data = self._create_letter_pdf(letter)
            documents.append((docx_data, pdf_data))
        
        return documents
    
    def _create_cv_docx(self, cv_data: Dict) -> bytes:
        """Create a DOCX version of the CV."""
        doc = Document()
        style = self.style_guide['cv_style']
        
        # Apply styles from style guide
        for section in doc.sections:
            section.left_margin = Inches(style['margins']['left'])
            section.right_margin = Inches(style['margins']['right'])
            section.top_margin = Inches(style['margins']['top'])
            section.bottom_margin = Inches(style['margins']['bottom'])
        
        # Add content (implementation would depend on cv_data structure)
        # This is a placeholder for the actual implementation
        doc.add_heading('CV Title', 0)
        
        # Save to bytes
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.read()
    
    def _create_cv_pdf(self, cv_data: Dict) -> bytes:
        """Create a PDF version of the CV using FPDF."""
        pdf = FPDF()
        style = self.style_guide['cv_style']
        
        pdf.add_page()
        pdf.set_margins(
            style['margins']['left'] * 25.4,
            style['margins']['top'] * 25.4,
            style['margins']['right'] * 25.4
        )
        
        # Add content
        pdf.set_font(style['font']['main'], size=style['font']['size'])
        pdf.multi_cell(0, style['spacing']['line_spacing'] * 10, cv_data['content'])
        
        # Save to bytes
        import io
        pdf_bytes = io.BytesIO()
        pdf.output(pdf_bytes)
        pdf_bytes.seek(0)
        return pdf_bytes.read()
    
    def _create_letter_docx(self, letter: Dict) -> bytes:
        """Create a DOCX version of a cover letter."""
        doc = Document()
        style = self.style_guide['cover_letter_style']
        
        # Apply styles from style guide
        for section in doc.sections:
            section.left_margin = Inches(style['margins']['left'])
            section.right_margin = Inches(style['margins']['right'])
            section.top_margin = Inches(style['margins']['top'])
            section.bottom_margin = Inches(style['margins']['bottom'])
        
        # Add content
        doc.add_paragraph(letter['content'])
        
        # Save to bytes
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.read()
    
    def _create_letter_pdf(self, letter: Dict) -> bytes:
        """Create a PDF version of a cover letter."""
        pdf = FPDF()
        style = self.style_guide['cover_letter_style']
        
        pdf.add_page()
        pdf.set_margins(
            style['margins']['left'] * 25.4,
            style['margins']['top'] * 25.4,
            style['margins']['right'] * 25.4
        )
        
        # Add content
        pdf.set_font(style['font']['main'], size=style['font']['size'])
        pdf.multi_cell(0, style['spacing']['line_spacing'] * 10, letter['content'])
        
        # Save to bytes
        import io
        pdf_bytes = io.BytesIO()
        pdf.output(pdf_bytes)
        pdf_bytes.seek(0)
        return pdf_bytes.read()